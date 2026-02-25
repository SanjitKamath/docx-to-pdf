from docx import Document
from docx.shared import Inches, RGBColor
import requests
import io

def create_complex_docx(filename):
    print("Generating complex test document...")
    doc = Document()

    # 1. Header (Tests margins and recurring elements)
    header = doc.sections[0].header
    header.paragraphs[0].text = "CONFIDENTIAL - SYSTEM STRESS TEST"

    # 2. Title and Subtitle
    doc.add_heading('DOCX to PDF Conversion Stress Test', 0)
    
    # 3. Inline Formatting (Bold, Italics, Colors)
    p = doc.add_paragraph('This document is designed to test the layout engine of ')
    p.add_run('LibreOffice').bold = True
    p.add_run(' inside a Docker container. It contains ')
    p.add_run('various formatting types,').italic = True
    
    color_run = p.add_run(' colored text,')
    color_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Red
    p.add_run(' and complex structures.')

    # 4. Lists (Tests indentation and bullet rendering)
    doc.add_heading('1. Lists and Indentation', level=1)
    doc.add_paragraph('First item in an unordered list', style='List Bullet')
    doc.add_paragraph('Second item in an unordered list', style='List Bullet')
    doc.add_paragraph('Step one of a process', style='List Number')
    doc.add_paragraph('Step two of a process', style='List Number')

    # 5. Tables (Tests grid layout and cell borders)
    doc.add_heading('2. Complex Table', level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Project Name'
    hdr_cells[2].text = 'Budget'
    
    # Data Rows
    table.rows[1].cells[0].text = '001'
    table.rows[1].cells[1].text = 'Alpha Pipeline'
    table.rows[1].cells[2].text = '$12,500.00'
    
    table.rows[2].cells[0].text = '002'
    table.rows[2].cells[1].text = 'Bravo Deployment'
    table.rows[2].cells[2].text = '$8,400.00'

    # 6. Embedded Images (Tests binary embedding and scaling)
    doc.add_heading('3. Embedded Image', level=1)
    doc.add_paragraph('Below is an image fetched from the web to test rendering:')
    try:
        # Download a placeholder image into memory
        image_url = "https://via.placeholder.com/400x150.png?text=PDF+Image+Test"
        response = requests.get(image_url)
        image_stream = io.BytesIO(response.content)
        
        # Insert it into the document
        doc.add_picture(image_stream, width=Inches(4.0))
    except Exception as e:
        doc.add_paragraph(f"[Image download failed: {e}]")

    # 7. Pagination (Tests page breaks and multi-page rendering)
    doc.add_page_break()
    doc.add_heading('4. Pagination Test', level=1)
    doc.add_paragraph('This text is explicitly placed on Page 2 to ensure page breaks and headers carry over correctly in the final PDF.')

    # Save the file
    doc.save(filename)
    print(f"Success! Saved as {filename}")
